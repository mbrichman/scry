#!/usr/bin/env python3
"""
DOCX Structure Probe

Analyzes Word documents to determine conversation structure and parsing strategy.
Generates metrics to help calibrate the hybrid parsing approach.

Usage:
    python scripts/utils/docx_probe.py sampe_word_docs/*.docx
    python scripts/utils/docx_probe.py sampe_word_docs/*.docx --json
"""

import sys
import json
import re
from pathlib import Path
from typing import Dict, List, Any, Tuple
from docx import Document


# Expanded role marker patterns
ROLE_MARKERS = {
    'user': [
        'You:', 'You said:', 'User:', 'Me:', 'Human:', 'Q:', 'Question:', 'Prompt:',
        '[You]', '[User]', '[Human]', '**You**', '**User**'
    ],
    'assistant': [
        'ChatGPT:', 'ChatGPT said:', 'Assistant:', 'AI:', 'A:', 'Answer:', 'Response:',
        'Claude:', 'Claude said:', 'GPT:', '[Assistant]', '[AI]', '[ChatGPT]',
        '**ChatGPT**', '**Assistant**', '**Claude**', '**AI**'
    ],
    'system': [
        'System:', '[System]', '**System**'
    ]
}

# Separator patterns
SEPARATORS = [
    r'^-{3,}$',  # ---
    r'^_{3,}$',  # ___
    r'^#{3,}$',  # ###
    r'^={3,}$',  # ===
    r'^\*{3,}$', # ***
]

# False positive patterns (things that look like markers but aren't)
FALSE_POSITIVES = [
    r'^[A-Z]\.',           # List items: A. B. C.
    r'^Note:',             # Notes
    r'^Example:',          # Examples
    r'^Step \d+:',         # Step 1:, Step 2:
    r'^\d+\.',             # Numbered lists: 1. 2. 3.
    r'^```',               # Code blocks
    r'^\s{4,}',            # Indented code
]


def clean_text(text: str) -> str:
    """Clean text by removing extra whitespace."""
    text = text.replace('\xa0', ' ').replace('\u00a0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def is_false_positive(text: str) -> bool:
    """Check if text matches false positive patterns."""
    for pattern in FALSE_POSITIVES:
        if re.match(pattern, text.strip()):
            return True
    return False


def detect_role_marker(text: str) -> Tuple[str | None, str]:
    """
    Detect if text starts with a role marker.
    
    Returns:
        Tuple of (role, marker_text) or (None, '') if no marker found
    """
    text_stripped = text.strip()
    
    # Check for false positives first
    if is_false_positive(text_stripped):
        return None, ''
    
    for role, markers in ROLE_MARKERS.items():
        for marker in markers:
            # Case-insensitive check
            if text_stripped.lower().startswith(marker.lower()):
                return role, marker
    
    return None, ''


def is_separator(text: str) -> bool:
    """Check if text is a separator line."""
    text_stripped = text.strip()
    for pattern in SEPARATORS:
        if re.match(pattern, text_stripped):
            return True
    return False


def is_likely_heading(paragraph) -> bool:
    """Heuristic to detect if paragraph is a heading."""
    if not paragraph.text.strip():
        return False
    
    # Check if paragraph has heading style
    if paragraph.style.name.startswith('Heading'):
        return True
    
    # Check if text is short and bold
    if len(paragraph.text.strip()) < 100:
        if paragraph.runs:
            # Check if all runs are bold
            all_bold = all(run.bold for run in paragraph.runs if run.text.strip())
            if all_bold:
                return True
    
    return False


def has_code_like_content(text: str) -> bool:
    """Detect if text contains code-like content."""
    indicators = [
        '```',              # Markdown code fence
        'def ',             # Python function
        'function ',        # JavaScript function
        'class ',           # Class definition
        '    ',             # Indentation (4 spaces)
        '\t',               # Tab indentation
        '{',                # Braces
        '};',               # Statement terminator
        'import ',          # Import statement
        '#include',         # C include
    ]
    
    return any(ind in text for ind in indicators)


def analyze_document(file_path: Path) -> Dict[str, Any]:
    """
    Analyze a DOCX file and generate structure metrics.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Dictionary of metrics and analysis results
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        return {
            'file': str(file_path),
            'error': str(e),
            'status': 'failed'
        }
    
    # Extract paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    non_empty = [p for p in paragraphs if p]
    
    # Count markers
    marker_counts = {'user': 0, 'assistant': 0, 'system': 0, 'total': 0}
    marked_lines = []
    consecutive_markers = []
    separator_count = 0
    heading_count = 0
    code_like_count = 0
    longest_unmarked_run = 0
    current_unmarked_run = 0
    
    prev_role = None
    
    for idx, (para_obj, text) in enumerate(zip(doc.paragraphs, paragraphs)):
        if not text:
            continue
        
        # Check for separator
        if is_separator(text):
            separator_count += 1
            continue
        
        # Check for heading
        if is_likely_heading(para_obj):
            heading_count += 1
        
        # Check for code-like content
        if has_code_like_content(text):
            code_like_count += 1
        
        # Check for role marker
        role, marker = detect_role_marker(text)
        
        if role:
            marker_counts[role] += 1
            marker_counts['total'] += 1
            
            # Record sample marked line
            if len(marked_lines) < 5:
                marked_lines.append({
                    'index': idx,
                    'role': role,
                    'marker': marker,
                    'preview': text[:80] + ('...' if len(text) > 80 else '')
                })
            
            # Track alternation
            if prev_role and prev_role != role:
                consecutive_markers.append((prev_role, role))
            
            prev_role = role
            
            # Reset unmarked run
            if current_unmarked_run > longest_unmarked_run:
                longest_unmarked_run = current_unmarked_run
            current_unmarked_run = 0
        else:
            current_unmarked_run += 1
    
    # Final unmarked run check
    if current_unmarked_run > longest_unmarked_run:
        longest_unmarked_run = current_unmarked_run
    
    # Calculate metrics
    total_paragraphs = len(paragraphs)
    non_empty_count = len(non_empty)
    
    coverage_ratio = marker_counts['total'] / non_empty_count if non_empty_count > 0 else 0
    
    # Alternation ratio: how often do markers alternate between user/assistant
    alternation_count = len([
        (a, b) for a, b in consecutive_markers 
        if {a, b} == {'user', 'assistant'}
    ])
    alternation_ratio = alternation_count / len(consecutive_markers) if consecutive_markers else 0
    
    # Determine suggested strategy
    # High coverage + good alternation = structured
    # Otherwise = semantic
    confidence_score = (coverage_ratio * 0.7) + (alternation_ratio * 0.3)
    suggested_strategy = 'structured' if confidence_score >= 0.4 else 'semantic'
    
    return {
        'file': file_path.name,
        'file_path': str(file_path),
        'status': 'success',
        'paragraphs': {
            'total': total_paragraphs,
            'non_empty': non_empty_count,
            'empty': total_paragraphs - non_empty_count,
        },
        'markers': {
            'user': marker_counts['user'],
            'assistant': marker_counts['assistant'],
            'system': marker_counts['system'],
            'total': marker_counts['total'],
            'coverage_ratio': round(coverage_ratio, 3),
        },
        'structure': {
            'separator_count': separator_count,
            'heading_count': heading_count,
            'code_like_count': code_like_count,
            'longest_unmarked_run': longest_unmarked_run,
        },
        'alternation': {
            'consecutive_pairs': len(consecutive_markers),
            'user_assistant_alternations': alternation_count,
            'alternation_ratio': round(alternation_ratio, 3),
        },
        'scoring': {
            'confidence_score': round(confidence_score, 3),
            'suggested_strategy': suggested_strategy,
        },
        'samples': {
            'marked_lines': marked_lines,
        },
        'flags': {
            'has_tables': len(doc.tables) > 0,
            'table_count': len(doc.tables),
            'has_code_blocks': code_like_count > 0,
        }
    }


def print_report(results: List[Dict[str, Any]], as_json: bool = False):
    """Print analysis report."""
    if as_json:
        print(json.dumps(results, indent=2))
        return
    
    print("=" * 80)
    print("DOCX STRUCTURE ANALYSIS REPORT")
    print("=" * 80)
    print()
    
    for result in results:
        if result.get('status') == 'failed':
            print(f"❌ {result['file']}")
            print(f"   Error: {result['error']}")
            print()
            continue
        
        print(f"📄 {result['file']}")
        print(f"   Path: {result['file_path']}")
        print()
        
        # Paragraphs
        print(f"   Paragraphs: {result['paragraphs']['total']} total, "
              f"{result['paragraphs']['non_empty']} non-empty")
        
        # Markers
        markers = result['markers']
        print(f"   Markers: {markers['total']} total "
              f"(user: {markers['user']}, assistant: {markers['assistant']}, "
              f"system: {markers['system']})")
        print(f"   Coverage: {markers['coverage_ratio']:.1%}")
        
        # Alternation
        alt = result['alternation']
        print(f"   Alternation: {alt['user_assistant_alternations']}/{alt['consecutive_pairs']} pairs "
              f"({alt['alternation_ratio']:.1%})")
        
        # Structure
        struct = result['structure']
        print(f"   Structure: {struct['separator_count']} separators, "
              f"{struct['heading_count']} headings, "
              f"{struct['code_like_count']} code blocks")
        print(f"   Longest unmarked run: {struct['longest_unmarked_run']} paragraphs")
        
        # Scoring
        score = result['scoring']
        strategy_emoji = "📋" if score['suggested_strategy'] == 'structured' else "📦"
        print(f"   {strategy_emoji} Confidence: {score['confidence_score']:.1%} → "
              f"Strategy: {score['suggested_strategy'].upper()}")
        
        # Flags
        flags = result['flags']
        if flags['has_tables']:
            print(f"   ⚠️  Contains {flags['table_count']} table(s)")
        if flags['has_code_blocks']:
            print(f"   💻 Contains code-like content")
        
        # Sample markers
        if result['samples']['marked_lines']:
            print(f"   Sample markers:")
            for sample in result['samples']['marked_lines'][:3]:
                print(f"      [{sample['index']}] {sample['role']:10s} → {sample['preview']}")
        
        print()
        print("-" * 80)
        print()
    
    # Summary
    print()
    print("SUMMARY")
    print("-" * 80)
    total_files = len(results)
    successful = len([r for r in results if r.get('status') == 'success'])
    failed = total_files - successful
    
    print(f"Total files: {total_files}")
    print(f"Successful: {successful}")
    print(f"Failed: {failed}")
    
    if successful > 0:
        structured = len([r for r in results if r.get('scoring', {}).get('suggested_strategy') == 'structured'])
        semantic = successful - structured
        
        print(f"\nSuggested strategies:")
        print(f"  Structured: {structured} ({structured/successful:.1%})")
        print(f"  Semantic: {semantic} ({semantic/successful:.1%})")
        
        avg_coverage = sum(r['markers']['coverage_ratio'] for r in results if r.get('status') == 'success') / successful
        avg_alternation = sum(r['alternation']['alternation_ratio'] for r in results if r.get('status') == 'success') / successful
        
        print(f"\nAverage metrics:")
        print(f"  Coverage ratio: {avg_coverage:.1%}")
        print(f"  Alternation ratio: {avg_alternation:.1%}")


def main():
    """Main entry point."""
    if len(sys.argv) < 2:
        print("Usage: python scripts/utils/docx_probe.py <file1.docx> [file2.docx] ...")
        print("       python scripts/utils/docx_probe.py sampe_word_docs/*.docx")
        print("       python scripts/utils/docx_probe.py sampe_word_docs/*.docx --json")
        sys.exit(1)
    
    # Parse arguments
    files = []
    as_json = False
    
    for arg in sys.argv[1:]:
        if arg == '--json':
            as_json = True
        else:
            files.append(Path(arg))
    
    if not files:
        print("Error: No files specified")
        sys.exit(1)
    
    # Analyze files
    results = []
    for file_path in files:
        if not file_path.exists():
            results.append({
                'file': str(file_path),
                'error': 'File not found',
                'status': 'failed'
            })
            continue
        
        result = analyze_document(file_path)
        results.append(result)
    
    # Print report
    print_report(results, as_json=as_json)


if __name__ == '__main__':
    main()
