#!/usr/bin/env python3
"""
Convert Claude's extracted content (markdown-like) to proper Word documents
"""
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def markdown_to_docx(text_content, output_path):
    """Convert markdown-like text to Word document with formatting"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = text_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue
            
        # Handle small caps headers like **[Career History]{.smallcaps}**
        if re.match(r'\*\*\[.*?\]\{\.smallcaps\}\*\*', line):
            text = re.sub(r'\*\*\[(.*?)\]\{\.smallcaps\}\*\*', r'\1', line)
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        # Handle company names with small caps like **i[Manage]{.smallcaps} LLC**
        elif re.match(r'\*\*.*?\{\.smallcaps\}.*?\*\*', line):
            text = re.sub(r'\*\*([^\[]*)\[([^\]]*)\]\{\.smallcaps\}([^\*]*)\*\*', r'\1\2\3', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove any remaining bold
            p = doc.add_heading(text, level=2)
            
        # Handle job titles like **Senior Director Product Management (January 2024--Present):**
        elif re.match(r'\*\*[^:]*:.*?\*\*', line):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            
        # Handle company descriptions in italics like *Document Management market leader...*
        elif re.match(r'^\*[^*].*\*$', line):
            text = re.sub(r'^\*(.*)\*$', r'\1', line)
            p = doc.add_paragraph(text)
            p.runs[0].italic = True
            
        # Handle bullet points
        elif line.startswith('-   ') or line.startswith('•   '):
            text = re.sub(r'^[-•]\s+', '', line)
            p = doc.add_paragraph(text, style='List Bullet')
            
        # Handle location/contact info with bullets
        elif '•' in line and not line.startswith('-'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Handle page headers like "Page 2"
        elif re.match(r'^Page \d+$', line):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Handle regular paragraphs
        else:
            # Clean up any remaining markdown
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
            
            if text:
                doc.add_paragraph(text)
    
    doc.save(output_path)
    return output_path

# Test with the resume example
resume_text = """**[Mark Richman]{.smallcaps}**

Chicago, Illinois • 312.203.6261 • mark@markrichman.net •
https://www.linkedin.com/in/mbrichman

Transformative product and technology leader with 20+ years in
LegalTech, specializing in cloud-based collaborative SaaS solutions.
Proven track record of translating product vision into market-leading
platforms, driving business growth through strategic product development
and technical innovation

Product Management • Agile Software Development • SaaS • Team Leadership
& Motivation

Application and Cloud Architecture • Collaboration Solutions • Business
Development • Recruiting • Strategic Planning

**[Career History]{.smallcaps}**

**i[Manage]{.smallcaps} LLC** • *Chicago, Illinois* August 2019--Present

*Document Management market leader for professional services
organizations*

**Senior Director Product Management (January 2024--Present):** Building
an industry leading SaaS-based document management and collaboration
platform for the professional services industry

-   Built and led cloud platform product team, hiring and mentoring five
    product managers to oversee key platform functions including API &
    Content Services, Tools, Security, and Core Infrastructure. Scaled
    the team within the larger product management organization and
    established a cohesive strategy to drive platform innovation and
    operational excellence

-   Led top-level company OKR to drive operational efficiency and
    improve gross margins. Successfully reduced operating costs and
    increased gross margin to 68% by standardizing on Azure Kubernetes
    Service (AKS) across the cloud platform. Achieved SLA goal of 99.95%
    platform availability, effectively balancing cost reduction with
    platform reliability

-   Proposed and received approval for a new product initiative,
    projected to generate more than $15 million in net new revenue"""

if __name__ == "__main__":
    # Read the full resume content
    with open("/Users/markrichman/projects/dovos/extracted_attachments/040_Resume Review Request_msg2_att0.txt", "r") as f:
        full_resume_text = f.read()
    
    output_file = "/Users/markrichman/projects/dovos/extracted_attachments/complete_resume.docx"
    markdown_to_docx(full_resume_text, output_file)
    print(f"Created complete Word document: {output_file}")