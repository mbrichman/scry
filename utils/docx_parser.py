"""
Utility for parsing Word documents containing chat conversations.

Uses a hybrid parsing strategy:
1. Analyze document structure to detect role markers and conversation patterns
2. Choose optimal parsing strategy (structured vs semantic)
3. Fall back gracefully when parsing fails

Adapted from the legacy ChromaDB implementation to work with PostgreSQL backend.
"""

import re
from datetime import datetime
from typing import List, Dict, Tuple, Optional
from docx import Document
from dataclasses import dataclass


# ===== Data structures =====

@dataclass
class ParseAnalysis:
    strategy: str  # 'structured' or 'semantic'
    confidence: float
    marker_coverage: float
    alternation_ratio: float
    marker_count: int
    total_non_empty: int


# ===== Utilities =====

def extract_timestamp(text: str) -> Optional[str]:
    """
    Try to extract timestamps from text using regex.
    
    Returns ISO format timestamp string if found, None otherwise.
    """
    # Look for common date formats
    date_patterns = [
        (r"(\d{4}-\d{2}-\d{2})", "%Y-%m-%d"),  # YYYY-MM-DD
        (r"(\d{2}/\d{2}/\d{4})", "%m/%d/%Y"),  # MM/DD/YYYY
        (r"(\w+ \d{1,2}, \d{4})", "%B %d, %Y"),  # Month DD, YYYY
    ]

    for pattern, fmt in date_patterns:
        matches = re.findall(pattern, text)
        if matches:
            try:
                dt = datetime.strptime(matches[0], fmt)
                return dt.isoformat()
            except ValueError:
                continue
    
    return None


def clean_text_content(text: str) -> str:
    """Clean up text content by removing extra whitespace."""
    # Clean up non-breaking spaces and other Unicode issues
    text = text.replace('\xa0', ' ').replace('\u00a0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


# Expanded role marker patterns
ROLE_MARKERS = {
    'user': [
        'You:', 'You said:', 'User:', 'Me:', 'Human:', 'Q:', 'Question:', 'Prompt:',
        '[You]', '[User]', '[Human]', '**You**', '**User**'
    ],
    'assistant': [
        'ChatGPT:', 'ChatGPT said:', 'Assistant:', 'AI:', 'A:', 'Answer:', 'Response:',
        'Claude:', 'Claude said:', 'GPT:', '[Assistant]', '[AI]', '[ChatGPT]',
        '**ChatGPT**', '**Assistant**', '**Claude**', '**AI**'
    ],
    'system': [
        'System:', '[System]', '**System**'
    ]
}

# False positive patterns (things that look like markers but aren't)
FALSE_POSITIVES = [
    r'^[A-Z]\.',           # List items: A. B. C.
    r'^Note:',             # Notes
    r'^Example:',          # Examples
    r'^Step \d+:',         # Step 1:, Step 2:
    r'^\d+\.',             # Numbered lists: 1. 2. 3.
]


def detect_role_marker(text: str) -> Tuple[Optional[str], str]:
    """
    Detect if text starts with a role marker.
    
    Returns:
        Tuple of (role, marker_text) or (None, '') if no marker found
    """
    text_stripped = text.strip()
    
    # Check for false positives first
    for pattern in FALSE_POSITIVES:
        if re.match(pattern, text_stripped):
            return None, ''
    
    for role, markers in ROLE_MARKERS.items():
        for marker in markers:
            # Case-insensitive check
            if text_stripped.lower().startswith(marker.lower()):
                return role, marker
    
    return None, ''


def is_question(text: str) -> bool:
    """Detect if text appears to be a question (heuristic-based)."""
    text = text.strip()
    # Check for question marks
    if text.endswith('?'):
        return True
    # Check for question words at the start
    question_words = ['what', 'why', 'how', 'when', 'where', 'who', 'which', 'can', 'could', 'would', 'should', 'do', 'does', 'did', 'is', 'are', 'was', 'were']
    first_word = text.split()[0].lower() if text.split() else ''
    if first_word in question_words:
        return True
    return False


# ===== Document Analysis =====

def analyze_document_structure(doc: Document) -> ParseAnalysis:
    """
    Analyze document structure to determine optimal parsing strategy.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        ParseAnalysis with strategy recommendation and confidence metrics
    """
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    non_empty = [p for p in paragraphs if p]
    
    if not non_empty:
        return ParseAnalysis(
            strategy='semantic',
            confidence=0.0,
            marker_coverage=0.0,
            alternation_ratio=0.0,
            marker_count=0,
            total_non_empty=0
        )
    
    # Count markers and track alternation
    marker_counts = {'user': 0, 'assistant': 0, 'system': 0, 'total': 0}
    consecutive_markers = []
    prev_role = None
    
    for text in non_empty:
        role, _ = detect_role_marker(text)
        
        if role:
            marker_counts[role] += 1
            marker_counts['total'] += 1
            
            # Track alternation between user/assistant
            if prev_role and prev_role != role:
                consecutive_markers.append((prev_role, role))
            prev_role = role
    
    # Calculate metrics
    marker_coverage = marker_counts['total'] / len(non_empty)
    
    # Alternation ratio: how often markers alternate user ↔ assistant
    alternation_count = len([
        (a, b) for a, b in consecutive_markers 
        if {a, b} == {'user', 'assistant'}
    ])
    alternation_ratio = alternation_count / len(consecutive_markers) if consecutive_markers else 0.0
    
    # Calculate confidence score
    # Weight alternation heavily (70%) since low coverage is expected for multi-paragraph messages
    # Boost coverage score to account for legitimate low coverage
    coverage_component = min(marker_coverage * 2.0, 1.0) * 0.3
    alternation_component = alternation_ratio * 0.7
    confidence = coverage_component + alternation_component
    
    # Determine strategy
    # Use structured if we have markers AND good alternation
    # Threshold: 0.4 confidence (can have ~10% coverage if perfect alternation)
    strategy = 'structured' if confidence >= 0.4 else 'semantic'
    
    return ParseAnalysis(
        strategy=strategy,
        confidence=confidence,
        marker_coverage=marker_coverage,
        alternation_ratio=alternation_ratio,
        marker_count=marker_counts['total'],
        total_non_empty=len(non_empty)
    )


# ===== Main Parsing Functions =====

def parse_docx_file(file_path: str, original_filename: Optional[str] = None) -> Tuple[List[Dict], List[str], str]:
    """
    Parse a Word document containing a chat conversation using hybrid strategy.
    
    Strategy selection:
    1. Analyze document structure (markers, alternation)
    2. Choose 'structured' if clear role markers present
    3. Fall back to 'semantic' chunking if markers absent/weak
    
    Args:
        file_path: Path to the .docx file
        original_filename: Original filename to use for title (optional)
        
    Returns:
        Tuple of (messages, timestamps, title) where:
        - messages: List of dicts with 'role', 'content', and optional 'metadata' keys
        - timestamps: List of ISO timestamp strings extracted from document
        - title: Document title (filename without extension)
    """
    doc = Document(file_path)
    
    # Step 1: Analyze document structure
    analysis = analyze_document_structure(doc)
    
    print(f"📊 Document analysis: {analysis.strategy.upper()} strategy ")
    print(f"   (confidence: {analysis.confidence:.1%}, ")
    print(f"   coverage: {analysis.marker_coverage:.1%}, ")
    print(f"   alternation: {analysis.alternation_ratio:.1%})")
    
    # Step 2: Parse using selected strategy
    if analysis.strategy == 'structured':
        messages, timestamps = _parse_structured(doc)
        print(f"   Structured parsing returned {len(messages)} messages")
    else:
        messages, timestamps = _parse_semantic(doc)
        print(f"   Semantic parsing returned {len(messages)} messages")
    
    # Debug: Log if no messages
    if not messages:
        print(f"   ⚠️  WARNING: No messages extracted from document!")
        print(f"   File: {file_path}")
        print(f"   Total paragraphs: {len(doc.paragraphs)}")
    
    # Add metadata about parsing to first message
    if messages:
        if 'metadata' not in messages[0]:
            messages[0]['metadata'] = {}
        messages[0]['metadata'].update({
            'parsing_strategy': analysis.strategy,
            'parsing_confidence': round(analysis.confidence, 3),
            'marker_coverage': round(analysis.marker_coverage, 3)
        })

    # Step 3: Extract title and metadata
    import os
    title = None
    
    # Try to get title from document core properties
    if doc.core_properties.title:
        title = doc.core_properties.title
    
    # Fallback to original filename if provided, then to actual file path
    if not title:
        if original_filename:
            title = os.path.splitext(os.path.basename(original_filename))[0]
        else:
            title = os.path.splitext(os.path.basename(file_path))[0]
    
    # Extract document creation/modification dates from core properties
    # These are more reliable than file system timestamps
    if doc.core_properties.created:
        doc_created = doc.core_properties.created.isoformat()
        # Add to timestamps for use by importer
        if doc_created not in timestamps:
            timestamps.insert(0, doc_created)
    
    if doc.core_properties.modified:
        doc_modified = doc.core_properties.modified.isoformat()
        if doc_modified not in timestamps:
            timestamps.append(doc_modified)

    return messages, timestamps, title


# ===== Parsing Strategies =====

def _parse_structured(doc: Document) -> Tuple[List[Dict], List[str]]:
    """
    Parse document using structured role markers.
    
    Handles formats like:
    - "You said:" / "ChatGPT said:"
    - "User:" / "Assistant:"
    - Missing first marker (accumulates pre-role content)
    
    Args:
        doc: python-docx Document object
        
    Returns:
        Tuple of (messages, timestamps)
    """
    messages = []
    timestamps = []
    current_role = None
    current_content = []
    pre_role_content = []  # Content before any role markers

    for p in doc.paragraphs:
        text = p.text.strip()
        text = clean_text_content(text)
        
        if not text:
            # Empty paragraph - save current message if any
            if current_role and current_content:
                messages.append({
                    'role': current_role,
                    'content': '\n'.join(current_content)
                })
                current_content = []
            continue

        # Extract potential timestamp
        ts = extract_timestamp(text)
        if ts:
            timestamps.append(ts)

        # Check if this is a role marker using new detection
        role, marker = detect_role_marker(text)
        
        if role:
            # If we have pre-role content, save it as user message
            if pre_role_content and not messages:
                messages.append({
                    'role': 'user',
                    'content': '\n'.join(pre_role_content)
                })
                pre_role_content = []
            
            # Save previous content if any
            if current_role and current_content:
                messages.append({
                    'role': current_role,
                    'content': '\n'.join(current_content)
                })

            # Set new role
            current_role = role
            current_content = []
        else:
            # Add to current content or pre-role content
            cleaned_text = clean_text_content(text)
            if cleaned_text:
                if current_role:
                    current_content.append(cleaned_text)
                else:
                    # No role set yet, accumulate in pre-role content
                    pre_role_content.append(cleaned_text)

    # Don't forget the last message
    if current_role and current_content:
        messages.append({
            'role': current_role,
            'content': '\n'.join(current_content)
        })
    
    # If still no messages after structured parsing, something went wrong
    # This shouldn't happen if analysis chose 'structured'
    if not messages and pre_role_content:
        # Fallback: treat all content as single user message
        messages.append({
            'role': 'user',
            'content': '\n'.join(pre_role_content)
        })
    
    return messages, timestamps


def _parse_semantic(doc: Document) -> Tuple[List[Dict], List[str]]:
    """
    Parse document using semantic chunking (no explicit role markers).
    
    Strategies:
    1. Split on double line breaks (empty paragraphs)
    2. Group consecutive paragraphs into chunks
    3. Use heuristics to assign roles (questions vs statements)
    4. Assume alternating user/assistant pattern
    
    Args:
        doc: python-docx Document object
        
    Returns:
        Tuple of (messages, timestamps)
    """
    timestamps = []
    
    # Extract all paragraphs and timestamps
    all_paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            text = clean_text_content(text)
            all_paragraphs.append(text)
            
            # Extract timestamps
            ts = extract_timestamp(text)
            if ts:
                timestamps.append(ts)
    
    if not all_paragraphs:
        return [], []
    
    # Group consecutive paragraphs into chunks based on empty lines
    # Since we already skipped empty paragraphs, we need a different approach
    # Use paragraph boundaries from the original document
    chunks = []
    current_chunk = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            text = clean_text_content(text)
            current_chunk.append(text)
        else:
            # Empty line - end current chunk
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
    
    # Don't forget the last chunk
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    # If we didn't get good chunks, use simple heuristics
    if not chunks:
        # Fallback: each paragraph is its own chunk
        chunks = all_paragraphs
    
    # Assign roles to chunks using improved heuristics
    messages = []
    for i, chunk in enumerate(chunks):
        if not chunk.strip():
            continue
        
        # Improved heuristics for role assignment:
        # 1. First chunk is usually user (conversation starts with user question)
        # 2. Questions are usually user
        # 3. Long responses are usually assistant
        # 4. Otherwise alternate strictly
        
        if i == 0:
            # First chunk - check if it's clearly a question
            if is_question(chunk) or len(chunk) < 200:
                role = 'user'
            else:
                # Long first chunk - could be assistant, but safer to assume user
                role = 'user'
        else:
            # For subsequent chunks, alternate strictly
            # This is more reliable than trying to guess
            prev_role = messages[-1]['role'] if messages else 'assistant'
            role = 'assistant' if prev_role == 'user' else 'user'
        
        messages.append({
            'role': role,
            'content': chunk
        })
    
    return messages, timestamps


def _parse_alternating_conversation(doc: Document) -> List[Dict]:
    """
    DEPRECATED: Use _parse_semantic() instead.
    
    Legacy function for parsing alternating conversation.
    Kept for backward compatibility but not used by parse_docx_file.
    """
    messages = []
    paragraphs = [clean_text_content(p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    
    if not paragraphs:
        return []
    
    # Group consecutive non-empty paragraphs into message blocks
    # Empty lines separate messages
    current_block = []
    blocks = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            current_block.append(clean_text_content(text))
        else:
            if current_block:
                blocks.append('\n'.join(current_block))
                current_block = []
    
    # Don't forget the last block
    if current_block:
        blocks.append('\n'.join(current_block))
    
    if not blocks:
        return []
    
    # Determine roles: assume conversation starts with user question
    # Then alternates: user, assistant, user, assistant, etc.
    for i, block in enumerate(blocks):
        # Check if this looks like a question (likely user)
        # Or if it's at an even index (assuming user starts)
        if i % 2 == 0:
            # Even index = user turn
            role = 'user'
        else:
            # Odd index = assistant turn
            role = 'assistant'
        
        messages.append({
            'role': role,
            'content': block
        })
    
    return messages
