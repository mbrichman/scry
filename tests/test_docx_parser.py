"""
Test suite for DOCX parser to ensure correct role assignment and message order.
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
from utils.docx_parser import parse_docx_file, analyze_document_structure


class TestCanadianSnax:
    """Tests for Canadian Snax document (no role markers, semantic parsing)"""
    
    def test_document_structure_analysis(self):
        """Verify document is correctly identified as semantic"""
        from docx import Document
        doc = Document('sampe_word_docs/Canadian Snax.docx')
        
        analysis = analyze_document_structure(doc)
        
        assert analysis.strategy == 'semantic', "Should use semantic strategy (no markers)"
        assert analysis.marker_count == 0, "Should have zero role markers"
        assert analysis.confidence < 0.4, "Confidence should be low"
    
    def test_first_message_is_user(self):
        """First message should be user question"""
        messages, _, _ = parse_docx_file('sampe_word_docs/Canadian Snax.docx')
        
        assert len(messages) > 0, "Should have messages"
        assert messages[0]['role'] == 'user', "First message should be user"
        assert messages[0]['content'].startswith('What do you think'), \
            f"First message should start with 'What do you think', got: {messages[0]['content'][:50]}"
    
    def test_messages_alternate_roles(self):
        """Messages should strictly alternate between user and assistant"""
        messages, _, _ = parse_docx_file('sampe_word_docs/Canadian Snax.docx')
        
        for i in range(1, min(10, len(messages))):
            prev_role = messages[i-1]['role']
            curr_role = messages[i]['role']
            
            assert curr_role != prev_role, \
                f"Message {i} should alternate roles: prev={prev_role}, curr={curr_role}"
            assert curr_role in ['user', 'assistant'], \
                f"Message {i} has invalid role: {curr_role}"
    
    def test_message_order_preserved(self):
        """Messages should be in document order"""
        messages, _, _ = parse_docx_file('sampe_word_docs/Canadian Snax.docx')
        
        # Check specific known sequences from the document
        first_content = messages[0]['content']
        second_content = messages[1]['content']
        
        # First should be the question about the take
        assert 'What do you think of this take' in first_content
        
        # Second should be the response about it being valid/entertaining
        assert 'valid' in second_content or 'entertaining' in second_content
    
    def test_total_message_count(self):
        """Should have reasonable message count"""
        messages, _, _ = parse_docx_file('sampe_word_docs/Canadian Snax.docx')
        
        # Canadian Snax has ~157 chunks based on empty lines
        # Should be around that number, give or take
        assert 100 <= len(messages) <= 200, \
            f"Expected ~157 messages, got {len(messages)}"


class TestMarkAndPolitics:
    """Tests for Mark and Politics document (has role markers, structured parsing)"""
    
    def test_document_structure_analysis(self):
        """Verify document is correctly identified as structured"""
        from docx import Document
        doc = Document('sampe_word_docs/Mark and Politics.docx')
        
        analysis = analyze_document_structure(doc)
        
        assert analysis.strategy == 'structured', "Should use structured strategy"
        assert analysis.marker_count > 0, "Should have role markers"
        assert analysis.confidence >= 0.4, "Confidence should be high"
    
    def test_first_message_is_user(self):
        """First message should be user"""
        messages, _, _ = parse_docx_file('sampe_word_docs/Mark and Politics.docx')
        
        assert len(messages) > 0, "Should have messages"
        assert messages[0]['role'] == 'user', "First message should be user"
        assert 'degree' in messages[0]['content'].lower() or \
               'political science' in messages[0]['content'].lower()
    
    def test_markers_correctly_parsed(self):
        """Role markers should be correctly identified"""
        messages, _, _ = parse_docx_file('sampe_word_docs/Mark and Politics.docx')
        
        # Check alternation for first 10 messages
        for i in range(1, min(10, len(messages))):
            prev_role = messages[i-1]['role']
            curr_role = messages[i]['role']
            
            # Should alternate
            assert curr_role != prev_role, \
                f"Structured parsing failed: messages don't alternate at position {i}"


class TestChappellRoan:
    """Tests for Chappell Roan document (has role markers, structured parsing)"""
    
    def test_document_structure_analysis(self):
        """Verify document is correctly identified as structured"""
        from docx import Document
        doc = Document('sampe_word_docs/Chappell Roan & Feelings.docx')
        
        analysis = analyze_document_structure(doc)
        
        assert analysis.strategy == 'structured', "Should use structured strategy"
        assert analysis.confidence >= 0.4, "Should have high confidence"


def test_all_sample_docs_parse_without_error():
    """Ensure all sample documents can be parsed without crashing"""
    from pathlib import Path
    
    sample_docs = Path('sampe_word_docs').glob('*.docx')
    
    for doc_path in sample_docs:
        # Skip temporary Word files (start with ~$)
        if doc_path.name.startswith('~$'):
            continue
            
        messages, timestamps, title = parse_docx_file(str(doc_path))
        
        assert len(messages) > 0, f"{doc_path.name} has no messages"
        assert title, f"{doc_path.name} has no title"
        
        # All messages should have role and content
        for i, msg in enumerate(messages):
            assert 'role' in msg, f"{doc_path.name} message {i} missing role"
            assert 'content' in msg, f"{doc_path.name} message {i} missing content"
            assert msg['role'] in ['user', 'assistant', 'system'], \
                f"{doc_path.name} message {i} has invalid role: {msg['role']}"


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
